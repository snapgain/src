# -*- coding: utf-8 -*-
"""Renderizador DOCX do guia."""

from util import dedent_body
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CLAY = RGBColor(0xC1, 0x5F, 0x3C)
INK = RGBColor(0x1F, 0x1E, 0x1D)
SLATE = RGBColor(0x5A, 0x56, 0x51)
BOX_BG = "F4EFE9"
RULE_BG = "E8E0D6"

_ORDER_TBLPR = ["tblStyle", "tblpPr", "tblOverlap", "bidiVisual", "tblStyleRowBandSize",
                "tblStyleColBandSize", "tblW", "jc", "tblCellSpacing", "tblInd",
                "tblBorders", "shd", "tblLayout", "tblCellMar", "tblLook",
                "tblCaption", "tblDescription", "tblPrChange"]
_ORDER_TCPR = ["cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge", "tcBorders", "shd",
               "noWrap", "tcMar", "textDirection", "tcFitText", "vAlign", "hideMark",
               "headers", "cellIns", "cellDel", "cellMerge", "tcPrChange"]
_ORDER_PPR = ["pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr",
              "widowControl", "numPr", "suppressLineNumbers", "pBdr", "shd", "tabs",
              "suppressAutoHyphens", "kinsoku", "wordWrap", "overflowPunct",
              "topLinePunct", "autoSpaceDE", "autoSpaceDN", "bidi", "adjustRightInd",
              "snapToGrid", "spacing", "ind", "contextualSpacing", "mirrorIndents",
              "suppressOverlap", "jc", "textDirection", "textAlignment",
              "textboxTightWrap", "outlineLvl", "divId", "cnfStyle", "rPr", "sectPr",
              "pPrChange"]


def _local(tag):
    return tag.split("}")[-1]


def _insert_ordered(parent, child, order):
    idx = order.index(_local(child.tag))
    for existing in parent:
        name = _local(existing.tag)
        if name in order and order.index(name) > idx:
            existing.addprevious(child)
            return
    parent.append(child)


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    _insert_ordered(tcPr, shd, _ORDER_TCPR)


def _no_borders(table):
    tblPr = table._tbl.tblPr
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    _insert_ordered(tblPr, borders, _ORDER_TBLPR)


def _left_bar(cell, hexcolor="C15F3C"):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), hexcolor)
    borders.append(left)
    _insert_ordered(tcPr, borders, _ORDER_TCPR)


class DocxEmitter:
    def __init__(self):
        doc = Document()
        self.doc = doc
        sec = doc.sections[0]
        sec.top_margin = Cm(2.2)
        sec.bottom_margin = Cm(2.2)
        sec.left_margin = Cm(2.4)
        sec.right_margin = Cm(2.4)

        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = INK
        normal.paragraph_format.space_after = Pt(7)
        normal.paragraph_format.line_spacing = 1.18
        rfonts = normal.element.get_or_add_rPr().get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), "Calibri")

    # ---------------------------------------------------------- capa
    def capa(self, linha1, linha2, subtitulo, destaques, rodape):
        doc = self.doc
        for _ in range(4):
            doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run(linha1)
        r.font.size = Pt(40)
        r.font.bold = True
        r.font.color.rgb = INK
        p.paragraph_format.space_after = Pt(0)

        p = doc.add_paragraph()
        r = p.add_run(linha2)
        r.font.size = Pt(40)
        r.font.bold = True
        r.font.color.rgb = CLAY
        p.paragraph_format.space_after = Pt(14)

        p = doc.add_paragraph()
        r = p.add_run(subtitulo)
        r.font.size = Pt(13)
        r.font.color.rgb = SLATE
        p.paragraph_format.space_after = Pt(28)

        t = doc.add_table(rows=1, cols=1)
        _no_borders(t)
        c = t.cell(0, 0)
        _shade(c, BOX_BG)
        _left_bar(c)
        first = True
        for texto, bold in destaques:
            pp = c.paragraphs[0] if first else c.add_paragraph()
            first = False
            pp.paragraph_format.left_indent = Cm(0.3)
            pp.paragraph_format.space_after = Pt(3)
            rr = pp.add_run("•  " + texto)
            rr.font.size = Pt(11)
            rr.bold = bold
            rr.font.color.rgb = INK
        c.paragraphs[0].paragraph_format.space_before = Pt(9)
        c.paragraphs[-1].paragraph_format.space_after = Pt(10)

        for _ in range(6):
            doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run(rodape)
        r.font.size = Pt(9.5)
        r.font.color.rgb = SLATE

    # ------------------------------------------------------ títulos
    def h1(self, text, page_break=True):
        doc = self.doc
        if page_break:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text.upper())
        r.font.size = Pt(19)
        r.font.bold = True
        r.font.color.rgb = CLAY
        rule = doc.add_paragraph()
        rule.paragraph_format.space_after = Pt(12)
        pPr = rule._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "8")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "C15F3C")
        pbdr.append(bot)
        _insert_ordered(pPr, pbdr, _ORDER_PPR)

    def h2(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = INK

    def h3(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.color.rgb = CLAY

    # -------------------------------------------------------- texto
    def _runs(self, p, text, size=10.5, color=INK, italic=False):
        for i, part in enumerate(text.split("**")):
            if not part:
                continue
            r = p.add_run(part)
            r.font.size = Pt(size)
            r.font.color.rgb = color
            r.italic = italic
            if i % 2 == 1:
                r.bold = True

    def para(self, text, italic=False, muted=False, after=7):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after)
        self._runs(p, text, color=SLATE if muted else INK, italic=italic)

    def bullet(self, text, level=0):
        p = self.doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.7 + 0.6 * level)
        p.paragraph_format.space_after = Pt(3)
        self._runs(p, text)

    def numbered(self, text):
        p = self.doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.space_after = Pt(3)
        self._runs(p, text)

    # ------------------------------------------------------- blocos
    def prompt(self, titulo, corpo, nota=None):
        t = self.doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        _no_borders(t)
        cell = t.cell(0, 0)
        _shade(cell, BOX_BG)
        _left_bar(cell)
        cell.width = Cm(16.2)

        p0 = cell.paragraphs[0]
        p0.paragraph_format.space_before = Pt(6)
        p0.paragraph_format.space_after = Pt(4)
        p0.paragraph_format.left_indent = Cm(0.25)
        r0 = p0.add_run("PROMPT · " + titulo.upper())
        r0.font.size = Pt(8.5)
        r0.font.bold = True
        r0.font.color.rgb = CLAY

        for linha in dedent_body(corpo).split("\n"):
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(0.25)
            p.paragraph_format.line_spacing = 1.12
            r = p.add_run(linha)
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            r.font.color.rgb = INK

        if nota:
            pn = cell.add_paragraph()
            pn.paragraph_format.space_before = Pt(5)
            pn.paragraph_format.space_after = Pt(6)
            pn.paragraph_format.left_indent = Cm(0.25)
            rn = pn.add_run("↳ " + nota)
            rn.font.size = Pt(8.5)
            rn.italic = True
            rn.font.color.rgb = SLATE
        else:
            cell.paragraphs[-1].paragraph_format.space_after = Pt(8)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def tabela(self, headers, rows, widths=None):
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = t.rows[0].cells
        for i, htxt in enumerate(headers):
            _shade(hdr[i], "2B2926")
            p = hdr[i].paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(htxt)
            r.font.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for ri, row in enumerate(rows):
            cells = t.add_row().cells
            for i, val in enumerate(row):
                if ri % 2 == 1:
                    _shade(cells[i], "F7F4F0")
                p = cells[i].paragraphs[0]
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                self._runs(p, val, size=9.5)
        if widths:
            for i, w in enumerate(widths):
                for row in t.rows:
                    row.cells[i].width = Cm(w)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def callout(self, titulo, texto):
        t = self.doc.add_table(rows=1, cols=1)
        _no_borders(t)
        cell = t.cell(0, 0)
        _shade(cell, RULE_BG)
        p0 = cell.paragraphs[0]
        p0.paragraph_format.space_before = Pt(6)
        p0.paragraph_format.space_after = Pt(2)
        p0.paragraph_format.left_indent = Cm(0.25)
        r0 = p0.add_run(titulo)
        r0.font.size = Pt(9.5)
        r0.font.bold = True
        r0.font.color.rgb = INK
        p1 = cell.add_paragraph()
        p1.paragraph_format.space_after = Pt(7)
        p1.paragraph_format.left_indent = Cm(0.25)
        r1 = p1.add_run(texto)
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = INK
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def save(self, path):
        self.doc.save(path)
